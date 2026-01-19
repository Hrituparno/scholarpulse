import json
import logging
from pathlib import Path
from utils import now_iso, save_text, save_json

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generates research reports in Markdown, JSON, DOCX, and TXT formats."""

    def __init__(self, out_dir="output"):
        self.out_dir = Path(out_dir)

    def generate_report(self, query, papers, new_ideas, report_sections=None, experiment=None, results=None):
        ts = now_iso()
        # Clean timestamp for filenames (Windows safe)
        timestamp_clean = ts.replace(':', '-').replace('.', '-')
        
        # --- Markdown Generation (With Icons) ---
        md = [f"# 📑 ScholarPulse Research Report: {query.title()}\n"]
        md.append(f"**Generated**: {ts}\n\n")

        # Fallback if sections missing
        if not report_sections:
            report_sections = {
                "introduction": "Introduction generation failed.",
                "the_issue": "Issue details unavailable.",
                "conclusion": "Conclusion unavailable."
            }

        # I. INTRODUCTION
        md.append("## I. INTRODUCTION 📝\n")
        md.append(f"{report_sections.get('introduction', '')}\n\n")

        # II. THE ISSUE
        md.append("## II. THE ISSUE ⚠️\n")
        md.append(f"{report_sections.get('the_issue', '')}\n\n")

        # III. LITERATURE REVIEW
        md.append("## III. LITERATURE REVIEW 📑\n")
        md.append("The current state of the field is characterized by a rapid shift towards the following key breakthroughs. Below is a detailed synthesis of the most relevant academic contributions:\n\n")

        # Detailed List
        md.append("### Detailed Review\n")
        for i, p in enumerate(papers, 1):
            scholar_link = f"([Google Scholar]({p.get('google_scholar_url', '#')}))"
            md.append(f"#### {i}. {p['title']} {scholar_link}\n")
            md.append(f"**Authors**: {', '.join(p['authors'][:3])}\n\n")
            md.append(f"**Objective**: {p.get('objective', 'N/A')}\n\n")
            md.append(f"**Summary**: {p['summary']}\n\n")
            md.append(f"**Method**: {p.get('method', 'N/A')}\n")
            md.append(f"**Tools**: {p.get('tools', 'N/A')}\n")
            md.append(f"**Results**: {p.get('results', 'N/A')}\n")
            md.append("---\n")

        # IV. RECOMMENDATIONS
        md.append("## IV. RECOMMENDATIONS 💡\n")
        md.append("### Proposed New Research Directions\n")
        
        if isinstance(new_ideas, list):
            for i, idea in enumerate(new_ideas, 1):
                md.append(f"**{i}. {idea.get('title', 'Untitled')}** 🚀\n")
                md.append(f"{idea.get('description', '')}\n")
                reqs = idea.get('requirements', [])
                req_str = ", ".join(reqs) if isinstance(reqs, list) else str(reqs)
                md.append(f"*Relevance/Requirements*: {req_str}\n\n")
        else:
            md.append(str(new_ideas) + "\n")

        # V. CONCLUSION
        md.append("## V. CONCLUSION ✅\n")
        md.append(f"{report_sections.get('conclusion', '')}\n\n")

        # Save MD
        body = "\n".join(md)
        # Fix filename format: ensure no double extensions or invalid chars
        out_md = self.out_dir / f"report_{timestamp_clean}.md"
        out_json = self.out_dir / f"report_{timestamp_clean}.json"
        
        if not self.out_dir.exists():
            self.out_dir.mkdir(parents=True, exist_ok=True)
            
        save_text(body, out_md)
        
        # --- DOCX Generation ---
        if HAS_DOCX:
            try:
                docx_path = self.out_dir / f"report_{timestamp_clean}.docx"
                self._save_docx(docx_path, query, papers, new_ideas, report_sections, ts)
                print(f"[ScholarPulse] Word Report saved to: {docx_path}")
            except Exception as e:
                logger.error(f"Failed to generate DOCX: {e}")
                print(f"[ScholarPulse] Failed to generate DOCX: {e}")

        # --- TXT Generation ---
        txt_path = self.out_dir / f"report_{timestamp_clean}.txt"
        self._save_txt(txt_path, query, papers, new_ideas, report_sections)
        print(f"[ScholarPulse] Text Report saved to: {txt_path}")
        
        # Save JSON last
        save_json({
            "query": query, 
            "papers": papers, 
            "new_ideas": new_ideas, 
            "report_sections": report_sections, 
            "experiment": experiment, 
            "results": results, 
            "generated_at": ts
        }, out_json)

        return str(out_md)

    def _save_txt(self, path, query, papers, new_ideas, sections):
        """Generates a clean text file version."""
        lines = []
        lines.append(f"RESEARCH REPORT: {query.upper()}")
        lines.append("==================================================\n")
        
        lines.append("I. INTRODUCTION")
        lines.append(sections.get("introduction", ""))
        lines.append("\nII. THE ISSUE")
        lines.append(sections.get("the_issue", ""))
        
        lines.append("\nIII. LITERATURE REVIEW")
        lines.append("--------------------------------------------------")
        for i, p in enumerate(papers, 1):
            lines.append(f"{i}. {p['title']}")
            lines.append(f"   Objective: {p.get('objective', 'N/A')}")
            lines.append(f"   Method: {p.get('method', 'N/A')}")
            lines.append(f"   Results: {p.get('results', 'N/A')}")
            lines.append("")
            
        lines.append("IV. RECOMMENDATIONS")
        lines.append("--------------------------------------------------")
        if isinstance(new_ideas, list):
            for i, idea in enumerate(new_ideas, 1):
                lines.append(f"{i}. {idea.get('title', 'Untitled')}")
                lines.append(f"   {idea.get('description', '')}")
                lines.append("")
                
        lines.append("\nV. CONCLUSION")
        lines.append(sections.get("conclusion", ""))
        
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

    def _save_docx(self, path, query, papers, new_ideas, sections, ts):
        """Generates a beautified Docx file with tables."""
        doc = Document()
        
        # Title
        heading = doc.add_heading(f"Research Report: {query.title()}", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        p = doc.add_paragraph(f"Generated: {ts}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # I. Introduction
        doc.add_heading('I. INTRODUCTION', level=1)
        doc.add_paragraph(sections.get('introduction', ''))
        
        # II. The Issue
        doc.add_heading('II. THE ISSUE', level=1)
        doc.add_paragraph(sections.get('the_issue', ''))
        
        # III. Literature Review
        doc.add_heading('III. LITERATURE REVIEW', level=1)
        doc.add_paragraph("Analysis of the contemporary research landscape and critical assessment of key findings:")
        
        # Detailed Analysis
        doc.add_heading('Deep Technical Analysis', level=2)
        for i, p in enumerate(papers, 1):
            p_para = doc.add_paragraph()
            runner = p_para.add_run(f"{i}. {p['title']}")
            runner.bold = True
            runner.font.size = Pt(12)
            
            doc.add_paragraph(f"Domain: {p.get('application', 'Inferred from context')}")
            doc.add_paragraph(f"Methodology: {p.get('method', 'Domain standard')}")
            doc.add_paragraph(f"Industrial Tools: {p.get('tools', 'Standard infrastructure')}")
            doc.add_paragraph(f"Core Objective: {p.get('objective', 'Synthesis pending')}")
            
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run("Abstract Synthesis: ")
            summary_run.italic = True
            summary_para.add_run(p['summary'])
        
        # IV. Recommendations
        doc.add_heading('IV. RECOMMENDATIONS', level=1)
        if isinstance(new_ideas, list):
            for i, idea in enumerate(new_ideas, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{i}. {idea.get('title', 'Untitled')}")
                run.bold = True
                run.font.size = Pt(11)
                
                desc = doc.add_paragraph(idea.get('description', ''))
                desc.style = 'List Bullet'
                
                reqs = idea.get('requirements', [])
                req_str = ", ".join(reqs) if isinstance(reqs, list) else str(reqs)
                if req_str:
                    r = doc.add_paragraph(f"Requirements: {req_str}")
                    r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    r.runs[0].italic = True
        
        # V. Conclusion
        doc.add_heading('V. CONCLUSION', level=1)
        doc.add_paragraph(sections.get('conclusion', ''))
        
        doc.save(path)
